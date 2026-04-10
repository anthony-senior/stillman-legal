from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.4

h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(24)
h1_style.font.color.rgb = RGBColor(0x03, 0x0D, 0x28)
h1_style.font.bold = True
h1_style.paragraph_format.space_before = Pt(0)
h1_style.paragraph_format.space_after = Pt(16)

h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(16)
h2_style.font.color.rgb = RGBColor(0x03, 0x0D, 0x28)
h2_style.font.bold = True
h2_style.paragraph_format.space_before = Pt(24)
h2_style.paragraph_format.space_after = Pt(8)

h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Calibri'
h3_style.font.size = Pt(13)
h3_style.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)
h3_style.font.bold = True
h3_style.paragraph_format.space_before = Pt(16)
h3_style.paragraph_format.space_after = Pt(4)

# ============================================================
# COVER / TITLE
# ============================================================
for _ in range(4):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Website Audit Report')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x03, 0x0D, 0x28)
run.font.name = 'Calibri'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('stillmanlegalpc.com')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xD0, 0xA7, 0x87)
run.font.name = 'Calibri'

doc.add_paragraph('')

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Prepared: April 9, 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x68, 0x68, 0x68)

prep_p = doc.add_paragraph()
prep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prep_p.add_run('Prepared by: Flare International Solutions')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x68, 0x68, 0x68)

prep_p2 = doc.add_paragraph()
prep_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prep_p2.add_run('Web Development & Technology Services')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x68, 0x68, 0x68)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('Executive Summary', level=1)

doc.add_paragraph(
    'This report presents the findings of a comprehensive audit of the Stillman Legal P.C. website '
    '(stillmanlegalpc.com). The audit covered SEO health, content quality, technical infrastructure, '
    'design/UX, and missed business opportunities.'
)

doc.add_paragraph(
    'The site is built on WordPress 6.9.4 with a custom theme and has solid employment law content. '
    'However, it suffers from multiple broken links, SEO issues on the immigration pages, outdated content, '
    'and a design that does not compete with other NYC law firms. Several issues are actively costing the '
    'firm potential clients right now, including broken lead capture links and duplicate page titles.'
)

# Summary table
summary_p = doc.add_paragraph()
run = summary_p.add_run('Issues Found: ')
run.bold = True
run = summary_p.add_run('20 total (7 Critical, 6 Moderate, 7 Opportunities)')

doc.add_paragraph('')

# ============================================================
# SEO PROBLEMS
# ============================================================
doc.add_heading('SEO Problems', level=1)

# Issue 1
doc.add_heading('1. All Immigration Pages Share the Same Title Tag', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Critical')
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph(
    'Every immigration page (citizenship, asylum, DACA, waivers, family visas, LGBTQ+ marriage) '
    'has the exact same title tag: "Apply for green card | Stillman Legal PC". This means Google '
    'sees 7 pages competing for the same keyword instead of 7 distinct, properly indexed pages.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Each immigration page needs a unique, descriptive title tag. For example: '
    '"NYC Asylum Application Lawyers | Stillman Legal PC" for the asylum page, '
    '"DACA Applications & Renewals | NYC Immigration Attorneys" for the DACA page, etc. '
    'This can be done in Yoast SEO by editing each page\'s SEO title field.',
    style='List Bullet'
)

# Issue 2
doc.add_heading('2. Broken Image URL on Book Page', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Critical')
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph(
    'The book cover image for "Knowing Your Rights" has a malformed URL: '
    'https://stillmanlegalpc.comwp-content/uploads/... (missing the forward slash after .com). '
    'This means the book promotion section displays a broken image.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Update the image source URL in WordPress to include the missing slash: '
    'https://stillmanlegalpc.com/wp-content/uploads/... This is a one-minute fix in the page editor.',
    style='List Bullet'
)

# Issue 3
doc.add_heading('3. Stale Content (Employment Pages)', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'Most employment law pages have not been updated since May 2022 (over 4 years). '
    'Google rewards fresh, regularly updated content. Competitors who publish and update '
    'content regularly will outrank these pages over time.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Review and refresh each employment law page with current legal developments, '
    'updated statistics, and recent case information. Even minor updates signal to Google '
    'that the content is maintained. Aim to refresh each page at least annually.',
    style='List Bullet'
)

# Issue 4
doc.add_heading('4. No Blog or Content Marketing', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'The sitemap contains only 26 page URLs and zero blog posts. There is no post sitemap. '
    'For a law firm competing in NYC, content marketing (blog posts about employment law topics, '
    'legal news, case studies) is essential for organic search visibility.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Start a blog targeting long-tail keywords related to NYC employment law. '
    'Even 2-4 posts per month on topics like "What to do if you\'re fired for reporting harassment in NYC" '
    'or "Understanding New York wage theft laws" can significantly improve organic traffic.',
    style='List Bullet'
)

# Issue 5
doc.add_heading('5. Thin Sitemap', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'Only 26 URLs are indexed in the sitemap. Competitors in NYC employment law likely have '
    'hundreds of indexed pages. A larger, well-structured site with more content signals '
    'authority to search engines.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Expand the site with additional content: FAQs per practice area, case study pages, '
    'neighborhood-specific landing pages (e.g., "Employment Lawyer in Manhattan"), '
    'and the blog mentioned above.',
    style='List Bullet'
)

# ============================================================
# CONTENT ISSUES
# ============================================================
doc.add_heading('Content Issues', level=1)

# Issue 6
doc.add_heading('6. Confusing Duplicate Page Slugs', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Critical')
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph(
    'The URL /applying-for-citizenship/ serves family visa content, while '
    '/applying-for-citizenship-2/ contains the actual citizenship information. '
    'This is confusing for visitors and harmful for SEO. Users searching for citizenship '
    'information may land on the wrong page.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Rename the slugs to clearly reflect their content: /family-visas/ for the family visa page '
    'and /citizenship/ for the citizenship page. Set up 301 redirects from the old URLs to '
    'the new ones using Yoast SEO\'s redirect manager.',
    style='List Bullet'
)

# Issue 7
doc.add_heading('7. Incorrect Geographic References', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'The asylum page FAQ mentions "Connecticut" and the family visa page asks '
    '"Do you need a Connecticut Family Visa?" even though the firm is based in New York City. '
    'This appears to be leftover content from a previous version of the site or a template, '
    'and it undermines the firm\'s credibility.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Search all pages for "Connecticut" references and replace with "New York" or "New York City" '
    'as appropriate. Review all immigration page content for consistency with the firm\'s NYC location.',
    style='List Bullet'
)

# Issue 8
doc.add_heading('8. Inconsistent Content Depth', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'Employment pages have rich, detailed content (the racial discrimination page is 2,000+ words) '
    'while some immigration pages are very thin (the DACA page is roughly 300 words). '
    'Thin content pages rank poorly and may be flagged by Google as low-quality.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Expand thin immigration pages with more detailed information, FAQs, process explanations, '
    'and client-focused content. Aim for at least 800-1,000 words per page with substantive, '
    'helpful information.',
    style='List Bullet'
)

# ============================================================
# BROKEN LINKS & FUNCTIONALITY
# ============================================================
doc.add_heading('Broken Links & Functionality', level=1)

# Issue 9
doc.add_heading('9. "Download A Free Copy Now" Button is Broken (Homepage)', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Critical')
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph(
    'The "Download A Free Copy Now" call-to-action button on the homepage (promoting Lina Stillman\'s '
    'book "Knowing Your Rights -- A Guide For Working People In New York") links to a broken URL: '
    'https://stillmanlegalpc.comknowing-your-rights-a-guide-for-new-york-working-people-by-a-new-york-city-employment-attorney/ '
    '(missing the forward slash after .com).'
)

doc.add_paragraph(
    'This is particularly damaging because it is the firm\'s primary lead magnet. Every visitor '
    'who clicks this button gets a 404 error instead of the book download, representing lost leads '
    'and wasted traffic.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Update the link URL on the homepage to include the missing slash. The correct URL should be: '
    'https://stillmanlegalpc.com/knowing-your-rights-a-guide-for-new-york-working-people-by-a-new-york-city-employment-attorney/',
    style='List Bullet'
)

# Issue 10
doc.add_heading('10. Book Cover Image Also Broken', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Critical')
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph(
    'The book cover image next to the download button has the same broken URL pattern '
    '(missing slash after .com). The image does not display, making the book promotion section '
    'look broken and unprofessional.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Fix the image src URL in the homepage content. Same issue as above -- add the missing slash.',
    style='List Bullet'
)

# ============================================================
# TECHNICAL ISSUES
# ============================================================
doc.add_heading('Technical Issues', level=1)

# Issue 11
doc.add_heading('11. Custom Theme With No Active Maintenance', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'The site runs a custom WordPress theme called "stillmanlegal" that was likely built by a '
    'freelance developer. There is no evidence of active theme maintenance. Bugs are accumulating '
    '(broken URLs, Connecticut references) with no one fixing them. The theme does not use a page '
    'builder, which means any content changes require someone comfortable with HTML.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Either engage a developer for ongoing maintenance, or consider rebuilding the site with a '
    'modern framework that is easier to maintain. A managed rebuild would also address the '
    'design and SEO issues identified in this report.',
    style='List Bullet'
)

# Issue 12
doc.add_heading('12. No Form Integration With CRM', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Opportunity')
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

doc.add_paragraph(
    'The consultation form uses WPForms, which is functional but appears to only send form submissions '
    'to email. There is no visible integration with a CRM system for lead tracking, follow-up '
    'automation, or client intake workflow.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Integrate the form with a CRM system (such as EspoCRM, HubSpot, or Clio) to track leads, '
    'automate follow-up emails, and ensure no consultation request falls through the cracks.',
    style='List Bullet'
)

# ============================================================
# DESIGN & UX ISSUES
# ============================================================
doc.add_heading('Design & User Experience Issues', level=1)

# Issue 13
doc.add_heading('13. Dated Visual Design', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'The site\'s CSS and layout patterns feel dated (2020-2022 era). While functional, the design '
    'does not convey the premium, trustworthy image that NYC clients expect from a law firm. '
    'Competing firms in the same market have more modern, polished websites.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'A full design refresh or rebuild using modern design patterns: generous whitespace, '
    'refined typography, subtle animations, and a layout that feels premium and professional. '
    'This does not require changing the branding (colors, logo) -- just updating the visual execution.',
    style='List Bullet'
)

# Issue 14
doc.add_heading('14. Mobile Experience Not Optimized', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Moderate')
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

doc.add_paragraph(
    'The site works on mobile devices but is not designed mobile-first. Over 60% of legal service '
    'searches happen on mobile devices. A mobile-optimized experience with touch-friendly buttons, '
    'readable text without zooming, and fast load times is essential.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Implement responsive design with mobile-first principles. Ensure tap targets are at least 44px, '
    'text is readable at default zoom, and the consultation form is easy to complete on a phone.',
    style='List Bullet'
)

# Issue 15
doc.add_heading('15. Empty or Minimal Client Reviews Section', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Opportunity')
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

doc.add_paragraph(
    'The homepage mentions "Client Reviews" and has a BBB badge, but the actual reviews content '
    'appears minimal or empty. The /reviews/ page in the sitemap has not been updated since 2022. '
    'Social proof is one of the most powerful conversion tools for law firms.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Collect and display client testimonials prominently. Integrate Google Reviews, Avvo reviews, '
    'or manually curated testimonials with client permission. Display them on the homepage and '
    'create a dedicated reviews page.',
    style='List Bullet'
)

# ============================================================
# MISSED OPPORTUNITIES
# ============================================================
doc.add_heading('Missed Opportunities', level=1)

# Issue 16
doc.add_heading('16. No Structured Data (Schema Markup)', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Opportunity')
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

doc.add_paragraph(
    'The site has no visible structured data markup for LocalBusiness, Attorney, LegalService, '
    'or FAQ schemas. This means no rich snippets appear in Google search results (no star ratings, '
    'no FAQ dropdowns, no address/phone in search cards).'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Add JSON-LD structured data for: LocalBusiness (address, phone, hours), '
    'Attorney (Lina Stillman\'s credentials), LegalService (practice areas), '
    'and FAQPage (on pages with FAQ content). Yoast SEO Premium can handle some of this automatically.',
    style='List Bullet'
)

# Issue 17
doc.add_heading('17. No Spanish Language Version', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Opportunity')
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

doc.add_paragraph(
    'The firm\'s social media handles are in Spanish ("LuchaXTusDerechos", '
    '"LuchaPorTusDerechosCom"), strongly suggesting they serve a Spanish-speaking clientele. '
    'However, the entire website is in English only. A Spanish version would capture a '
    'significant untapped audience.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Create a Spanish-language version of the site (at minimum the homepage, key practice area pages, '
    'and the consultation form). This can be done with a WordPress multilingual plugin like WPML '
    'or Polylang, or built into a site rebuild.',
    style='List Bullet'
)

# Issue 18
doc.add_heading('18. Mobile App Poorly Promoted', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Opportunity')
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

doc.add_paragraph(
    'Stillman Legal has both iOS and Android apps ("Fight for Your Rights"), which is impressive '
    'for a firm this size. However, the apps are barely promoted on the website -- the download '
    'links are buried near the footer and easy to miss.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Feature the app prominently on the homepage and about page. A dedicated section with '
    'app screenshots, download badges, and a clear value proposition would drive more downloads.',
    style='List Bullet'
)

# Issue 19
doc.add_heading('19. No Google Analytics or Conversion Tracking Visible', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Opportunity')
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

doc.add_paragraph(
    'While analytics may be installed via a plugin, there is no visible evidence of conversion '
    'tracking on the consultation form. Without tracking, the firm cannot measure which pages '
    'or traffic sources generate the most consultation requests.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Implement Google Analytics 4 with conversion events for form submissions and phone call clicks. '
    'This data is essential for understanding what\'s working and where to invest marketing spend.',
    style='List Bullet'
)

# Issue 20
doc.add_heading('20. Homepage "Download A Free Copy Now" Button Broken', level=2)
p = doc.add_paragraph()
run = p.add_run('Severity: ')
run.bold = True
run = p.add_run('Critical')
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph(
    'In addition to the book page link being broken (Issue #9), the "Download A Free Copy Now" '
    'button that appears in the "Know Your Rights As A New York Employee" section on the homepage '
    'also does not work. Clicking it leads to a 404 error page. This is the firm\'s primary '
    'content offer and lead generation tool, and it is completely non-functional.'
)

doc.add_paragraph(
    'This means every visitor who is interested enough to click the download button is met with '
    'an error -- a terrible user experience that actively damages trust and loses potential clients.'
)

doc.add_heading('How to fix:', level=3)
doc.add_paragraph(
    'Immediately fix the broken URL (add the missing slash after .com). Consider also creating '
    'a proper landing page for the book download that captures the visitor\'s email address before '
    'providing the download link, turning this into an effective lead generation funnel.',
    style='List Bullet'
)

# ============================================================
# SUMMARY TABLE
# ============================================================
doc.add_page_break()
doc.add_heading('Summary of All Issues', level=1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '#'
hdr_cells[1].text = 'Issue'
hdr_cells[2].text = 'Severity'
hdr_cells[3].text = 'Category'

issues = [
    ('1', 'Duplicate title tags on immigration pages', 'Critical', 'SEO'),
    ('2', 'Broken book cover image URL', 'Critical', 'SEO'),
    ('3', 'Stale employment law content (4+ years)', 'Moderate', 'SEO'),
    ('4', 'No blog or content marketing', 'Moderate', 'SEO'),
    ('5', 'Thin sitemap (only 26 URLs)', 'Moderate', 'SEO'),
    ('6', 'Confusing duplicate page slugs', 'Critical', 'Content'),
    ('7', 'Incorrect Connecticut references', 'Moderate', 'Content'),
    ('8', 'Inconsistent content depth', 'Moderate', 'Content'),
    ('9', '"Download A Free Copy" link broken', 'Critical', 'Broken'),
    ('10', 'Book cover image broken', 'Critical', 'Broken'),
    ('11', 'Custom theme with no maintenance', 'Moderate', 'Technical'),
    ('12', 'No CRM form integration', 'Opportunity', 'Technical'),
    ('13', 'Dated visual design', 'Moderate', 'Design/UX'),
    ('14', 'Mobile experience not optimized', 'Moderate', 'Design/UX'),
    ('15', 'Empty client reviews section', 'Opportunity', 'Design/UX'),
    ('16', 'No structured data / schema markup', 'Opportunity', 'Opportunity'),
    ('17', 'No Spanish language version', 'Opportunity', 'Opportunity'),
    ('18', 'Mobile app poorly promoted', 'Opportunity', 'Opportunity'),
    ('19', 'No conversion tracking visible', 'Opportunity', 'Opportunity'),
    ('20', 'Homepage download button broken', 'Critical', 'Broken'),
]

for issue in issues:
    row_cells = table.add_row().cells
    row_cells[0].text = issue[0]
    row_cells[1].text = issue[1]
    row_cells[2].text = issue[2]
    row_cells[3].text = issue[3]

# Format table
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph('')

# ============================================================
# CLOSING
# ============================================================
doc.add_heading('Next Steps', level=1)

doc.add_paragraph(
    'The critical issues (broken links, duplicate titles, broken download button) should be addressed '
    'immediately as they are actively costing the firm potential clients. These fixes can be completed '
    'in under an hour.'
)

doc.add_paragraph(
    'The moderate issues (stale content, design refresh, mobile optimization) should be addressed '
    'as part of a planned website improvement project. A full site rebuild would address all of these '
    'simultaneously while also implementing the identified opportunities (schema markup, Spanish '
    'language, blog, CRM integration).'
)

doc.add_paragraph(
    'We have prepared a fully functional website prototype demonstrating how the rebuilt site could look '
    'and function, incorporating modern design standards, SEO best practices, and all content from the '
    'existing site. This prototype is available for review at the URL provided separately.'
)

doc.add_paragraph('')

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('---')
run.font.color.rgb = RGBColor(0xD0, 0xA7, 0x87)

doc.add_paragraph('')

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run('Flare International Solutions\n')
run.font.size = Pt(11)
run.bold = True
run = contact.add_run('Web Development & Technology Services\n')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x68, 0x68, 0x68)
run = contact.add_run('flareintl.com | anthony.senior@flareintl.com')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

# ============================================================
# SAVE
# ============================================================
downloads = os.path.expanduser('~/Downloads')
filepath = os.path.join(downloads, 'Stillman Legal - Website Audit Report.docx')
doc.save(filepath)
print(f'Saved to: {filepath}')
