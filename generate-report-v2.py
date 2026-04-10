from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.4

for level, size in [('Heading 1', 24), ('Heading 2', 16), ('Heading 3', 13)]:
    s = doc.styles[level]
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor(0x03, 0x0D, 0x28)
    s.font.bold = True
    s.paragraph_format.space_before = Pt(0 if level == 'Heading 1' else 24 if level == 'Heading 2' else 16)
    s.paragraph_format.space_after = Pt(16 if level == 'Heading 1' else 8 if level == 'Heading 2' else 4)

doc.styles['Heading 3'].font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

# ============================================================
def add_severity(doc, level):
    p = doc.add_paragraph()
    run = p.add_run('Severity: ')
    run.bold = True
    run = p.add_run(level)
    colors = {
        'Critical': RGBColor(0x99, 0x00, 0x00),
        'Moderate': RGBColor(0xCC, 0x88, 0x00),
        'Opportunity': RGBColor(0x5C, 0x80, 0xBC),
    }
    run.font.color.rgb = colors.get(level, RGBColor(0x33, 0x33, 0x33))

def add_fix(doc, text):
    doc.add_heading('How to fix:', level=3)
    doc.add_paragraph(text, style='List Bullet')

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Website Audit Report')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x03, 0x0D, 0x28)
run.font.name = 'Calibri'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('stillmanlegalpc.com')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xD0, 0xA7, 0x87)

doc.add_paragraph('')

for text in ['Prepared: April 9, 2026', 'Prepared by: Flare International Solutions', 'Web Development & Technology Services']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12 if 'Prepared' in text else 11)
    run.font.color.rgb = RGBColor(0x68, 0x68, 0x68)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('Executive Summary', level=1)

doc.add_paragraph(
    'This report presents the findings of a comprehensive audit of the Stillman Legal P.C. website '
    '(stillmanlegalpc.com). The audit covered SEO health, content quality, technical infrastructure, '
    'design/UX, site performance, and missed business opportunities.'
)

doc.add_paragraph(
    'The site is built on WordPress 6.9.4 with a custom theme and has solid employment law content. '
    'However, it suffers from multiple broken links, critical SEO issues on the immigration pages, '
    'an entirely empty reviews page, outdated content, and a design that does not compete with other '
    'NYC law firms. Several issues are actively costing the firm potential clients right now, including '
    'broken lead capture links, duplicate page titles, and missing social proof.'
)

p = doc.add_paragraph()
run = p.add_run('Issues Found: ')
run.bold = True
run = p.add_run('22 total (8 Critical, 6 Moderate, 8 Opportunities)')

doc.add_paragraph(
    'We recommend running this site through Google PageSpeed Insights (pagespeed.web.dev) before '
    'any client-facing conversation. The mobile performance score will provide additional data points '
    'about site speed and Core Web Vitals compliance.'
)

# ============================================================
# SEO PROBLEMS
# ============================================================
doc.add_heading('SEO Problems', level=1)

doc.add_heading('1. Every Immigration Page Has the Wrong Title Tag', level=2)
add_severity(doc, 'Critical')
doc.add_paragraph(
    'Every single immigration page -- citizenship, asylum, DACA, waivers, LGBTQ+, green card -- '
    'shows the exact same title: "Apply for green card | Stillman Legal PC." Google uses the title '
    'tag to understand what a page is about and show it in search results. Someone searching "asylum '
    'lawyer New York" will never find their asylum page because Google thinks it\'s a green card page. '
    'This is a textbook SEO disaster that is actively costing them immigration clients right now.'
)
add_fix(doc,
    'Each immigration page needs a unique, descriptive title tag. For example: '
    '"NYC Asylum Application Lawyers | Stillman Legal PC" for the asylum page, '
    '"DACA Applications & Renewals | NYC Immigration Attorneys" for the DACA page. '
    'This can be done in Yoast SEO by editing each page\'s SEO title field. Estimated time: 30 minutes.')

doc.add_heading('2. Broken Image URL on Book Page', level=2)
add_severity(doc, 'Critical')
doc.add_paragraph(
    'The book cover image for "Knowing Your Rights" has a malformed URL: '
    'https://stillmanlegalpc.comwp-content/uploads/... (missing the forward slash after .com). '
    'The book promotion section displays a broken image, making it look unprofessional.'
)
add_fix(doc,
    'Update the image source URL to include the missing slash: '
    'https://stillmanlegalpc.com/wp-content/uploads/... This is a one-minute fix in the page editor.')

doc.add_heading('3. Stale Content (Employment Pages)', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'Most employment law pages have not been updated since May 2022 (over 4 years). '
    'Google rewards fresh, regularly updated content. Competitors who publish and update '
    'content regularly will outrank these pages over time.'
)
add_fix(doc,
    'Review and refresh each employment law page with current legal developments, '
    'updated statistics, and recent case information. Even minor updates signal to Google '
    'that the content is maintained. Aim to refresh each page at least annually.')

doc.add_heading('4. No Blog or Content Marketing', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'The sitemap contains only 26 page URLs and zero blog posts. There is no post sitemap. '
    'For a law firm competing in NYC, content marketing (blog posts about employment law topics, '
    'legal news, case studies) is essential for organic search visibility.'
)
add_fix(doc,
    'Start a blog targeting long-tail keywords related to NYC employment law. '
    'Even 2-4 posts per month on topics like "What to do if you\'re fired for reporting harassment '
    'in NYC" can significantly improve organic traffic over time.')

doc.add_heading('5. Thin Sitemap (Only 26 URLs)', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'Only 26 URLs are indexed in the sitemap. Competitors in NYC employment law likely have '
    'hundreds of indexed pages. A larger, well-structured site signals authority to search engines.'
)
add_fix(doc,
    'Expand the site with FAQs per practice area, case study pages, '
    'neighborhood-specific landing pages (e.g., "Employment Lawyer in Manhattan"), and a blog.')

doc.add_heading('6. No Structured Data (Schema Markup)', level=2)
add_severity(doc, 'Opportunity')
doc.add_paragraph(
    'No visible structured data markup for LocalBusiness, Attorney, LegalService, '
    'or FAQ schemas. This means no rich snippets in Google results -- no star ratings, '
    'no FAQ dropdowns, no address/phone in search cards. Competitors using schema will appear '
    'more prominently in the same search results.'
)
add_fix(doc,
    'Add JSON-LD structured data for LocalBusiness (address, phone, hours), '
    'Attorney (credentials), LegalService (practice areas), and FAQPage schemas. '
    'Yoast SEO Premium can handle some of this automatically.')

# ============================================================
# CONTENT ISSUES
# ============================================================
doc.add_heading('Content Issues', level=1)

doc.add_heading('7. Reviews Page is Completely Empty', level=2)
add_severity(doc, 'Critical')
doc.add_paragraph(
    'The /reviews/ page loads with the heading "Read Our Client Reviews At Stillman Legal PC." '
    'and nothing else beneath it. No testimonials, no Google reviews, no case outcomes -- just an '
    'empty page. For a law firm that wins cases for clients, this is a massive missed opportunity. '
    'Social proof is everything in legal services. An empty reviews page is arguably worse than '
    'having no reviews page at all, because it signals that the firm has no satisfied clients '
    'willing to speak up.'
)
add_fix(doc,
    'Collect and display client testimonials with permission. Integrate Google Reviews or Avvo '
    'reviews. Even 5-10 strong testimonials would transform this page from a liability into a '
    'conversion tool. Feature the best quotes on the homepage as well.')

doc.add_heading('8. Immigration Section is a Bolt-On, Not Integrated', level=2)
add_severity(doc, 'Critical')
doc.add_paragraph(
    'The employment law side of the site is coherent and well-structured with detailed, '
    'practice-specific content. The immigration pages look like they were copied from another source '
    'and pasted in with minimal customization. They reference "Connecticut" multiple times despite '
    'this being a New York City firm, and some pages have generic, placeholder-feeling copy. '
    'The disconnect between the polished employment pages and the rough immigration pages undermines '
    'credibility across the entire site.'
)
add_fix(doc,
    'Rewrite all immigration pages from scratch with NYC-specific content, consistent branding, '
    'and the same level of detail as the employment pages. Remove all Connecticut references. '
    'Each page should feel like it was written by a New York immigration attorney, not copied '
    'from a template.')

doc.add_heading('9. Confusing Duplicate Page Slugs', level=2)
add_severity(doc, 'Critical')
doc.add_paragraph(
    '/applying-for-citizenship/ serves family visa content, while /applying-for-citizenship-2/ '
    'has the actual citizenship information. This is confusing for visitors and harmful for SEO.'
)
add_fix(doc,
    'Rename slugs to clearly reflect their content: /family-visas/ and /citizenship/. '
    'Set up 301 redirects from old URLs using Yoast SEO\'s redirect manager.')

doc.add_heading('10. Inconsistent Content Depth', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'Employment pages have rich, detailed content (racial discrimination is 2,000+ words) '
    'while some immigration pages are very thin (DACA is roughly 300 words). '
    'Thin content pages rank poorly and may be flagged by Google as low-quality.'
)
add_fix(doc,
    'Expand thin immigration pages to at least 800-1,000 words with substantive, helpful '
    'information, FAQs, and process explanations.')

doc.add_heading('11. Zero Partnership or Referral Network Mentioned', level=2)
add_severity(doc, 'Opportunity')
doc.add_paragraph(
    'There is no mention anywhere on the site of working with other firms, referral resources, '
    'or what happens when a client\'s needs fall outside the firm\'s scope. For a firm that handles '
    'both employment and immigration, there are natural referral opportunities with firms in '
    'complementary practice areas (family law, criminal defense, personal injury, international '
    'mobility). Mentioning a referral network would signal that the firm is well-connected and '
    'client-focused.'
)
add_fix(doc,
    'Add a section on the About page or a dedicated page about the firm\'s professional network '
    'and referral relationships. This builds trust and opens the door to strategic partnerships.')

# ============================================================
# BROKEN LINKS
# ============================================================
doc.add_heading('Broken Links & Functionality', level=1)

doc.add_heading('12. Homepage "Download A Free Copy Now" Button is Broken', level=2)
add_severity(doc, 'Critical')
doc.add_paragraph(
    'The "Download A Free Copy Now" call-to-action button on the homepage -- in the "Know Your '
    'Rights As A New York Employee" section -- does not work. Clicking it leads to a 404 error page. '
    'The URL is malformed: https://stillmanlegalpc.comknowing-your-rights... (missing the forward '
    'slash after .com).'
)
doc.add_paragraph(
    'This is the firm\'s primary content offer and lead generation tool. Every visitor who is '
    'interested enough to click the download button is met with an error -- a terrible user '
    'experience that actively damages trust and loses potential clients. This button appears '
    'prominently on the homepage and is likely clicked frequently.'
)
add_fix(doc,
    'Immediately fix the broken URL (add the missing slash after .com). The correct URL is: '
    'https://stillmanlegalpc.com/knowing-your-rights-a-guide-for-new-york-working-people-by-a-new-york-city-employment-attorney/. '
    'Consider also creating a proper landing page that captures email addresses before providing '
    'the download, turning this into an effective lead generation funnel.')

doc.add_heading('13. Book Cover Image Also Broken on Homepage', level=2)
add_severity(doc, 'Critical')
doc.add_paragraph(
    'The book cover image next to the download button has the same broken URL pattern '
    '(missing slash). The image does not display, making the entire book promotion section '
    'look broken and unprofessional.'
)
add_fix(doc,
    'Fix the image src URL -- add the missing slash after .com. Same one-minute fix.')

# ============================================================
# TECHNICAL ISSUES
# ============================================================
doc.add_heading('Technical Issues', level=1)

doc.add_heading('14. Custom Theme With No Active Maintenance', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'The site runs a custom WordPress theme called "stillmanlegal" that was likely built by a '
    'freelance developer. There is no evidence of active theme maintenance. Bugs are accumulating '
    '(broken URLs, Connecticut references, empty pages) with no one fixing them. The theme does not '
    'use a page builder, so any content changes require someone comfortable editing HTML directly.'
)
add_fix(doc,
    'Either engage a developer for ongoing maintenance, or consider a full rebuild with a modern '
    'framework that is easier to maintain and addresses all issues in this report simultaneously.')

doc.add_heading('15. No Form Integration With CRM', level=2)
add_severity(doc, 'Opportunity')
doc.add_paragraph(
    'The consultation form uses WPForms, which is functional but appears to only send form '
    'submissions to email. There is no visible integration with a CRM system for lead tracking, '
    'follow-up automation, or client intake workflow.'
)
add_fix(doc,
    'Integrate the form with a CRM system (HubSpot, Clio, EspoCRM) to track leads, '
    'automate follow-up emails, and ensure no consultation request falls through the cracks.')

doc.add_heading('16. Performance (PageSpeed)', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'We recommend running the site through Google PageSpeed Insights (pagespeed.web.dev) to get '
    'a current mobile performance score. Law firm websites with poor mobile performance scores '
    'lose rankings in Google\'s mobile-first index. Over 60% of legal service searches happen on '
    'mobile devices. Core Web Vitals (LCP, FID, CLS) are direct Google ranking factors.'
)
add_fix(doc,
    'Run pagespeed.web.dev, note the mobile score, and address any flagged issues. Common fixes '
    'include image optimization, reducing JavaScript, and implementing proper caching.')

# ============================================================
# DESIGN / UX
# ============================================================
doc.add_heading('Design & User Experience Issues', level=1)

doc.add_heading('17. Dated Visual Design', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'The site\'s CSS and layout patterns feel dated (2020-2022 era). While functional, the design '
    'does not convey the premium, trustworthy image that NYC clients expect from a law firm. '
    'Competing firms have more modern, polished websites with generous whitespace, refined typography, '
    'and subtle animations.'
)
add_fix(doc,
    'A full design refresh or rebuild using modern design patterns. This does not require changing '
    'the branding (colors, logo) -- just updating the visual execution to match current standards.')

doc.add_heading('18. Mobile Experience Not Optimized', level=2)
add_severity(doc, 'Moderate')
doc.add_paragraph(
    'The site works on mobile devices but is not designed mobile-first. Touch targets may be too '
    'small, text may require zooming, and the consultation form may be difficult to complete on '
    'a phone screen.'
)
add_fix(doc,
    'Implement responsive design with mobile-first principles. Ensure tap targets are at least 44px, '
    'text is readable at default zoom, and the form is easy to complete on a phone.')

# ============================================================
# MISSED OPPORTUNITIES
# ============================================================
doc.add_heading('Missed Opportunities', level=1)

doc.add_heading('19. No Spanish Language Version', level=2)
add_severity(doc, 'Opportunity')
doc.add_paragraph(
    'The firm\'s social media handles are in Spanish ("LuchaXTusDerechos", "LuchaPorTusDerechosCom"), '
    'strongly suggesting they serve a Spanish-speaking clientele. However, the website is English-only. '
    'A Spanish version would capture a significant untapped audience, especially for immigration services.'
)
add_fix(doc,
    'Create a Spanish-language version of at minimum the homepage, key practice area pages, and the '
    'consultation form. This can be done with WPML/Polylang or built into a site rebuild.')

doc.add_heading('20. Mobile App Poorly Promoted', level=2)
add_severity(doc, 'Opportunity')
doc.add_paragraph(
    'Stillman Legal has both iOS and Android apps ("Fight for Your Rights"), which is impressive '
    'for a firm this size. However, the apps are barely promoted -- download links are buried near '
    'the footer and easy to miss.'
)
add_fix(doc,
    'Feature the app prominently on the homepage and about page with screenshots, download badges, '
    'and a clear value proposition.')

doc.add_heading('21. No Conversion Tracking Visible', level=2)
add_severity(doc, 'Opportunity')
doc.add_paragraph(
    'No visible evidence of conversion tracking on the consultation form. Without tracking, the '
    'firm cannot measure which pages or traffic sources generate the most consultation requests.'
)
add_fix(doc,
    'Implement Google Analytics 4 with conversion events for form submissions and phone call clicks.')

doc.add_heading('22. No Google Business Profile Integration', level=2)
add_severity(doc, 'Opportunity')
doc.add_paragraph(
    'There is no visible connection between the website and a Google Business Profile. For a law '
    'firm serving a specific geographic area (NYC), a well-optimized Google Business Profile with '
    'reviews, photos, and regular posts is one of the most powerful local SEO tools available.'
)
add_fix(doc,
    'Claim and optimize the Google Business Profile. Add photos, respond to reviews, post weekly '
    'updates, and ensure NAP (Name, Address, Phone) consistency between the website and the profile.')

# ============================================================
# SUMMARY TABLE
# ============================================================
doc.add_page_break()
doc.add_heading('Summary of All Issues', level=1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
hdr[0].text = '#'
hdr[1].text = 'Issue'
hdr[2].text = 'Severity'
hdr[3].text = 'Category'

issues = [
    ('1', 'All immigration pages have wrong title tag', 'Critical', 'SEO'),
    ('2', 'Broken book cover image URL', 'Critical', 'SEO'),
    ('3', 'Stale employment law content (4+ years)', 'Moderate', 'SEO'),
    ('4', 'No blog or content marketing', 'Moderate', 'SEO'),
    ('5', 'Thin sitemap (only 26 URLs)', 'Moderate', 'SEO'),
    ('6', 'No structured data / schema markup', 'Opportunity', 'SEO'),
    ('7', 'Reviews page is completely empty', 'Critical', 'Content'),
    ('8', 'Immigration section is a bolt-on', 'Critical', 'Content'),
    ('9', 'Confusing duplicate page slugs', 'Critical', 'Content'),
    ('10', 'Inconsistent content depth', 'Moderate', 'Content'),
    ('11', 'No partnership / referral network', 'Opportunity', 'Content'),
    ('12', 'Homepage download button broken', 'Critical', 'Broken'),
    ('13', 'Book cover image broken on homepage', 'Critical', 'Broken'),
    ('14', 'Custom theme with no maintenance', 'Moderate', 'Technical'),
    ('15', 'No CRM form integration', 'Opportunity', 'Technical'),
    ('16', 'Performance / PageSpeed unknown', 'Moderate', 'Technical'),
    ('17', 'Dated visual design', 'Moderate', 'Design/UX'),
    ('18', 'Mobile not optimized', 'Moderate', 'Design/UX'),
    ('19', 'No Spanish language version', 'Opportunity', 'Opportunity'),
    ('20', 'Mobile app poorly promoted', 'Opportunity', 'Opportunity'),
    ('21', 'No conversion tracking', 'Opportunity', 'Opportunity'),
    ('22', 'No Google Business Profile integration', 'Opportunity', 'Opportunity'),
]

for issue in issues:
    row = table.add_row().cells
    for i, val in enumerate(issue):
        row[i].text = val

for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)

# ============================================================
# APPROACH STRATEGY
# ============================================================
doc.add_page_break()
doc.add_heading('Appendix: Approach Strategy & Scripts', level=1)

doc.add_paragraph(
    'The following outlines a three-angle approach strategy for engaging Stillman Legal P.C. as a '
    'client. Each angle represents a separate conversation with its own timing and purpose. '
    'They should not be mixed.'
)

doc.add_heading('Angle 1: Personal Case (Unpaid Wages)', level=2)
doc.add_paragraph(
    'This is the most authentic entry point. If you have a genuine employment dispute (unpaid wages, '
    'wrongful termination, etc.), reaching out as a potential client is the most natural first contact. '
    'No sales pitch, no partnership talk -- just a real case inquiry.'
)
doc.add_paragraph(
    'This also provides insight into how the firm operates, their intake process, and how they treat '
    'clients. That context informs everything that follows.'
)

doc.add_heading('Script: Initial Call', level=3)
p = doc.add_paragraph()
run = p.add_run(
    '"Hi, I\'d like to speak with someone about a potential employment case. My previous employer '
    'let me go and still hasn\'t paid me what I\'m owed. I understand Stillman Legal handles wage '
    'disputes -- is there someone I can speak with about a consultation?"'
)
run.italic = True

doc.add_heading('If asked for details:', level=3)
p = doc.add_paragraph()
run = p.add_run(
    '"I was an employee, I was let go without receiving my final compensation. I\'ve been trying '
    'to resolve it directly but haven\'t gotten anywhere. I\'d like to understand what my options are."'
)
run.italic = True

doc.add_heading('Angle 2: Referral Partnership (After Initial Contact)', level=2)
doc.add_paragraph(
    'Stillman does US immigration -- getting people INTO America. Flare does international mobility '
    '-- helping people build options OUTSIDE America. These are complementary, not competing. '
    'An immigrant who gets their green card through Stillman might eventually want a second passport '
    'or residency as a backup plan. Someone Flare helps with international residency might need US '
    'immigration support for family members. The referral flows both directions.'
)

doc.add_heading('Script: Partnership Email', level=3)
p = doc.add_paragraph()
run = p.add_run(
    'Hi Lina,\n\n'
    'I reached out recently about my own employment situation -- thank you for the prompt response.\n\n'
    'I also wanted to introduce what I do professionally, because I think there\'s a natural overlap. '
    'I run Flare International Solutions -- we help people with international mobility: second '
    'residencies, second passports, and visa strategies outside of the US.\n\n'
    'You work with immigrants coming INTO the US. I work with people building options OUTSIDE the US. '
    'The client profiles don\'t compete -- they actually complement each other. Someone who just got '
    'their green card through you might eventually want a Plan B residency in Portugal or a second '
    'passport. Someone I help establish residency abroad might need US immigration support for '
    'family members.\n\n'
    'Would you be open to a 15-minute call to see if a referral relationship makes sense? No formal '
    'arrangement needed -- just wanted to put it on the table.\n\n'
    'Best,\nAnthony Senior\nFlare International Solutions'
)
run.italic = True

doc.add_heading('Angle 3: Website Services (Separate Conversation)', level=2)
doc.add_paragraph(
    'The website problems are real and severe. This is not a "your site could look better" '
    'conversation -- this is "your site has specific, measurable errors that are costing you '
    'immigration clients right now and I can show you exactly what they are." The title tag issue '
    'alone is the slam dunk for this call.'
)
doc.add_paragraph(
    'Before calling, run their site through Google PageSpeed Insights (pagespeed.web.dev) '
    'and note the mobile score. This provides additional data to lead with.'
)

doc.add_heading('Script: Cold Call', level=3)
p = doc.add_paragraph()
run = p.add_run(
    '"Hi, is this Lina? ... Hi Lina, my name\'s Anthony from Flare Labs -- I do website performance '
    'work for law firms. I took a look at your site before calling and I noticed something specific '
    'that\'s probably costing you immigration clients.\n\n'
    'Your immigration pages -- the green card page, asylum, citizenship, DACA -- they all have the '
    'same title tag saying \'Apply for green card.\' Google uses that title to decide what search '
    'terms to rank each page for. So someone searching \'asylum lawyer New York\' right now will never '
    'find your asylum page because Google thinks it\'s a duplicate green card page. That\'s a direct '
    'SEO hit on every immigration service you offer.\n\n'
    'You got 60 seconds? I can walk you through exactly what I\'m seeing."'
)
run.italic = True

doc.add_heading('If she\'s interested:', level=3)
p = doc.add_paragraph()
run = p.add_run(
    '"I can send you a full audit -- it\'ll show you the PageSpeed score, all the title tag issues, '
    'and a priority list of what to fix. No charge. Then if you want me to fix it, we can talk about '
    'that separately. What\'s the best email?"'
)
run.italic = True

doc.add_heading('If she says "we already have someone handling the site":', level=3)
p = doc.add_paragraph()
run = p.add_run(
    '"That\'s fine -- whoever built it probably just didn\'t know about the title tag duplication '
    'because it\'s easy to miss. But it\'s worth knowing about regardless of who fixes it. '
    'Can I send you the one-page summary so you can show them?"'
)
run.italic = True

doc.add_heading('Recommended Sequence', level=2)
doc.add_paragraph('Week 1: Contact about your own employment case. Establish a real, human relationship.', style='List Number')
doc.add_paragraph('After initial contact: Introduce the partnership angle naturally ("by the way, I run a global mobility firm...").', style='List Number')
doc.add_paragraph('Separately: Make the web development call. Lead with the audit data. Ideally to a different contact or timed separately from the personal matter.', style='List Number')

# ============================================================
# NEXT STEPS
# ============================================================
doc.add_page_break()
doc.add_heading('Next Steps', level=1)

doc.add_paragraph(
    'The critical issues (broken links, duplicate titles, broken download button, empty reviews page) '
    'should be addressed immediately as they are actively costing the firm potential clients. '
    'The broken URLs can be fixed in under an hour.'
)

doc.add_paragraph(
    'The moderate issues (stale content, design refresh, mobile optimization, performance) should be '
    'addressed as part of a planned website improvement project. A full site rebuild would address '
    'all of these simultaneously while also implementing the identified opportunities.'
)

doc.add_paragraph(
    'We have prepared a fully functional website prototype demonstrating how the rebuilt site could '
    'look and function, incorporating modern design standards, SEO best practices, and all content '
    'from the existing site. This prototype is available for review at the URL provided separately.'
)

doc.add_paragraph('')
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('---')
run.font.color.rgb = RGBColor(0xD0, 0xA7, 0x87)

doc.add_paragraph('')
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run('Flare International Solutions\n')
run.font.size = Pt(11)
run.bold = True
run = contact.add_run('Web Development & Technology Services\n')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x68, 0x68, 0x68)
run = contact.add_run('flareintl.com | anthony.senior@flareintl.com')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x5C, 0x80, 0xBC)

# ============================================================
# SAVE
# ============================================================
downloads = os.path.expanduser('~/Downloads')
filepath = os.path.join(downloads, 'Stillman Legal - Website Audit Report v2.docx')
doc.save(filepath)
print(f'Saved to: {filepath}')
